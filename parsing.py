from pathlib import Path
from shutil import rmtree
from string import ascii_uppercase

from docx import Document
from docx.oxml.ns import qn

SYMBOL_TO_UNICODE = {
    # --- Greek Lowercase ---
    0xF061: "\u03b1",  # alpha
    0xF062: "\u03b2",  # beta
    0xF063: "\u03c7",  # chi
    0xF064: "\u03b4",  # delta
    0xF065: "\u03b5",  # epsilon
    0xF066: "\u03c6",  # phi
    0xF067: "\u03b3",  # gamma
    0xF068: "\u03b7",  # eta
    0xF069: "\u03b9",  # iota
    0xF06A: "\u03c6",  # phi (variation)
    0xF06B: "\u03ba",  # kappa
    0xF06C: "\u03bb",  # lambda
    0xF06D: "\u03bc",  # mu
    0xF06E: "\u03bd",  # nu
    0xF06F: "\u03bf",  # omicron
    0xF070: "\u03c0",  # pi
    0xF071: "\u03b8",  # theta
    0xF072: "\u03c1",  # rho
    0xF073: "\u03c3",  # sigma
    0xF074: "\u03c4",  # tau
    0xF075: "\u03c5",  # upsilon
    0xF076: "\u03d6",  # varpi
    0xF077: "\u03c9",  # omega
    0xF078: "\u03be",  # xi
    0xF079: "\u03c8",  # psi
    0xF07A: "\u03b6",  # zeta
    # --- Greek Uppercase ---
    0xF041: "\u0391",  # Alpha
    0xF042: "\u0392",  # Beta
    0xF043: "\u03a7",  # Chi
    0xF044: "\u0394",  # Delta
    0xF045: "\u0395",  # Epsilon
    0xF046: "\u03a6",  # Phi
    0xF047: "\u0393",  # Gamma
    0xF048: "\u0397",  # Eta
    0xF049: "\u0399",  # Iota
    0xF04B: "\u039a",  # Kappa
    0xF04C: "\u039b",  # Lambda
    0xF04D: "\u039c",  # Mu
    0xF04E: "\u039d",  # Nu
    0xF04F: "\u039f",  # Omicron
    0xF050: "\u03a0",  # Pi
    0xF051: "\u0398",  # Theta
    0xF052: "\u03a1",  # Rho
    0xF053: "\u03a3",  # Sigma
    0xF054: "\u03a4",  # Tau
    0xF055: "\u03a5",  # Upsilon
    0xF056: "\u03c2",  # sigma (final)
    0xF057: "\u03a9",  # Omega
    0xF058: "\u039e",  # Xi
    0xF059: "\u03a8",  # Psi
    0xF05A: "\u0396",  # Zeta
    # --- Common Math/Logic ---
    0xF022: "\u2200",  # forall
    0xF024: "\u2203",  # exists
    0xF02A: "\u2217",  # asterisk operator
    0xF02B: "\u002b",  # plus
    0xF02C: "\u002c",  # comma
    0xF02D: "\u2212",  # minus sign
    0xF02E: "\u002e",  # period
    0xF02F: "\u002f",  # fraction slash
    0xF030: "\u0030",  # 0
    0xF031: "\u0031",  # 1
    0xF0B0: "\u00b0",  # degree sign
    0xF0B1: "\u00b1",  # plus-minus
    0xF0B2: "\u2033",  # double prime
    0xF0B3: "\u2265",  # greater equal
    0xF0B4: "\u00d7",  # multiply
    0xF0B5: "\u221d",  # proportional to
    0xF0B6: "\u2202",  # partial diff
    0xF0B7: "\u2022",  # bullet
    0xF0B8: "\u00f7",  # divide
    0xF0B9: "\u2260",  # not equal
    0xF0BA: "\u2261",  # identical to
    0xF0BB: "\u2248",  # approx equal
    0xF0BC: "\u2026",  # ellipsis
    0xF0BD: "\u23d0",  # vertical arrow extension
    0xF0BE: "\u21b5",  # return arrow
    0xF0BF: "\u2135",  # aleph
    0xF0A3: "\u2264",  # less equal
    0xF0A5: "\u221e",  # infinity
    0xF0A6: "\u0192",  # function (florin)
    0xF0A7: "\u2663",  # club suit
    0xF0A8: "\u2666",  # diamond suit
    0xF0A9: "\u2665",  # heart suit
    0xF0AA: "\u2660",  # spade suit
    # Arrows
    0xF0AB: "\u2194",  # left-right arrow
    0xF0AC: "\u2190",  # left arrow
    0xF0AD: "\u2191",  # up arrow
    0xF0AE: "\u2192",  # right arrow
    0xF0AF: "\u2193",  # down arrow
}

NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
NS_WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"


class ExamQuestion:
    def __init__(self, q_id, content, images=None):
        self.id = q_id
        self.content = content
        self.images = images or []
        self.answers = []
        self.weight = "1"
        self.order = None

    def add_answer(self, label, is_correct, text, images=None):
        self.answers.append(
            {
                "label": label,
                "is_correct": is_correct == "1",
                "text": text,
                "images": images or [],
            }
        )

    def __repr__(self):
        return f"Q{self.id} W{self.weight}: {self.content[:40]}... ({len(self.answers)} ans)"


def get_formatted_text(cell):
    full_text = []

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run_text = ""
            for child in run._element:
                if child.tag == f"{NS_W}t":
                    run_text += child.text or ""
                elif child.tag == f"{NS_W}sym":
                    char_hex = child.get(f"{NS_W}char")
                    if char_hex:
                        code = int(char_hex, 16)
                        decoded = SYMBOL_TO_UNICODE.get(code)
                        if decoded:
                            run_text += decoded
                        else:
                            run_text += f"[SYM:{char_hex}]"
                elif child.tag == f"{NS_W}tab":
                    run_text += "\t"
                elif child.tag == f"{NS_W}br":
                    run_text += "\n"

            if not run_text:
                continue

            if run.font.superscript:
                full_text.append(f"<sup>{run_text}</sup>")
            elif run.font.subscript:
                full_text.append(f"<sub>{run_text}</sub>")
            else:
                full_text.append(run_text)

        full_text.append(" ")

    return "".join(full_text).strip()


def _extract_images_from_cell(cell, doc, image_output_dir, prefix):
    image_infos = []

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            for blip in run._element.iter(f"{NS_A}blip"):
                r_id = blip.get(qn("r:embed"))
                if not r_id:
                    continue

                image_part = doc.part.related_parts.get(r_id)
                if image_part is None:
                    continue

                inline_or_anchor = None
                parent = blip
                while parent is not None:
                    if parent.tag in (f"{NS_WP}inline", f"{NS_WP}anchor"):
                        inline_or_anchor = parent
                        break
                    parent = parent.getparent()

                width_emu = height_emu = None
                if inline_or_anchor is not None:
                    extent = inline_or_anchor.find(f"{NS_WP}extent")
                    if extent is not None:
                        cx = extent.get("cx")
                        cy = extent.get("cy")
                        if cx and cy:
                            try:
                                width_emu = int(cx)
                                height_emu = int(cy)
                            except ValueError:
                                width_emu = height_emu = None

                ext = Path(image_part.filename).suffix or ".png"
                filename = f"{prefix}_{len(image_infos) + 1}{ext}"
                out_path = image_output_dir / filename

                if not out_path.exists():
                    with open(out_path, "wb") as f:
                        f.write(image_part.blob)

                image_infos.append(
                    {
                        "path": str(out_path),
                        "width_emu": width_emu,
                        "height_emu": height_emu,
                    }
                )

    return image_infos


def parse_exam_file(file_path, output_dir):
    images_dir = Path(output_dir) / "Images"
    images_dir.mkdir()
    doc = Document(file_path)
    questions = []
    current_q = None

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells

            if len(cells) != 5:
                continue

            col_text = [cells[i].text.strip() for i in range(4)]

            if col_text[0].isdigit():
                q_id = col_text[0]
                content_text = get_formatted_text(cells[4])
                question_images = _extract_images_from_cell(
                    cells[4], doc, images_dir, f"q{q_id}"
                )
                current_q = ExamQuestion(q_id, content_text, question_images)
                questions.append(current_q)
            elif current_q and col_text[1] in ascii_uppercase:
                label = col_text[1]
                content_text = get_formatted_text(cells[4])
                answer_images = _extract_images_from_cell(
                    cells[4],
                    doc,
                    images_dir,
                    f"q{current_q.id}_ans_{label}",
                )
                current_q.add_answer(label, col_text[2], content_text, answer_images)
                if col_text[3].isdigit():
                    current_q.weight = col_text[3]

    rmtree(images_dir)
    return questions
