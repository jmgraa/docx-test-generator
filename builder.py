import re
from pathlib import Path
from string import ascii_uppercase

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Emu, Pt

FONT_NAME = "Calibri"
NORMAL_SIZE = Pt(12)
SMALL_SIZE = Pt(8)
SCRIPT_PATTERN = re.compile(
    r"(<sup>.*?</sup>|<sub>.*?</sub>|<eq>.*?</eq>)", re.IGNORECASE | re.DOTALL
)


def create_document(header_text, intro):
    doc = Document()

    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    run_h = header_para.add_run(header_text)
    run_h.font.size = SMALL_SIZE
    run_h.font.name = FONT_NAME

    _add_paragraphs_to_document(doc, intro.split("\n"))

    return doc


def add_question_on_document(doc, index, question, answers):
    p_q = doc.add_paragraph()
    _add_formatted_text(p_q, f"{index}) {question.content}")
    _add_picture(question.images, doc)

    for i, answer in enumerate(answers):
        label = ascii_uppercase[i]

        p_a = doc.add_paragraph()
        _add_formatted_text(p_a, f"{label}) {answer['text']}")
        _add_picture(answer["images"], doc)

    doc.add_paragraph()

    return doc


def _add_picture(images, doc):
    for image_info in images:
        img_path = image_info.get("path")
        width_emu = image_info.get("width_emu")
        height_emu = image_info.get("height_emu")

        if not img_path:
            continue

        p_a_img = doc.add_paragraph()
        run_a_img = p_a_img.add_run()
        try:
            if width_emu and height_emu:
                run_a_img.add_picture(
                    str(img_path),
                    width=Emu(width_emu),
                    height=Emu(height_emu),
                )
            else:
                run_a_img.add_picture(str(img_path))
        except Exception:
            p_a_img.add_run(f"[OBRAZEK: {Path(img_path).name}]")


def save_document(doc, index, dir):
    path = Path(dir) / f"{index}.docx"
    doc.save(path)


def save_data_file(prefix, extension, data, index, dir):
    path = Path(dir) / f"{prefix}{index}.{extension}"
    with open(path, "w") as f:
        f.write("\n".join(data))


def _add_paragraphs_to_document(doc, paragraphs):
    for text in paragraphs:
        p = doc.add_paragraph()
        _add_formatted_text(p, text)


def _extract_equations_with_placeholders(text):
    eq_pattern = re.compile(r"<eq>(.*?)</eq>", re.IGNORECASE | re.DOTALL)
    equations = []

    def replace_eq(match):
        equations.append(match.group(1))
        return f"__EQUATION_{len(equations)-1}__"

    return eq_pattern.sub(replace_eq, text), equations


def _add_formatted_text(paragraph, text):
    eq_placeholder_pattern = re.compile(r"__EQUATION_(\d+)__")
    text_with_placeholders, equations = _extract_equations_with_placeholders(text)
    parts = SCRIPT_PATTERN.split(text_with_placeholders)

    for part in parts:
        if not part:
            continue

        part_lower = part.lower()
        is_superscript = part_lower.startswith("<sup>") and part_lower.endswith(
            "</sup>"
        )
        is_subscript = part_lower.startswith("<sub>") and part_lower.endswith("</sub>")
        content = part[5:-6] if is_superscript or is_subscript else part

        if eq_placeholder_pattern.search(content):
            _process_mixed_content(
                paragraph,
                content,
                equations,
                eq_placeholder_pattern,
                is_superscript,
                is_subscript,
            )
        else:
            run = paragraph.add_run(content)
            run.font.superscript = is_superscript
            run.font.subscript = is_subscript
            run.font.size = NORMAL_SIZE
            run.font.name = FONT_NAME


def _process_mixed_content(
    paragraph,
    content,
    equations,
    eq_placeholder_pattern,
    is_superscript,
    is_subscript,
):
    segments = eq_placeholder_pattern.split(content)

    for i, segment in enumerate(segments):
        if not segment:
            continue

        if i % 2:
            eq_index = int(segment)
            omml_xml = equations[eq_index]
            try:
                if "xmlns:m=" not in omml_xml and omml_xml.strip().startswith("<"):
                    replacer = ""
                    if omml_xml.strip().startswith("<m:oMath"):
                        replacer = ("<m:oMath",)
                    elif omml_xml.strip().startswith("<oMath"):
                        replacer = "<oMath"

                    if replacer:
                        omml_xml = omml_xml.replace(
                            replacer,
                            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"',
                            1,
                        )
                oMath_element = parse_xml(omml_xml)
                paragraph._p.append(oMath_element)
            except Exception as e:
                run = paragraph.add_run(f"[EQUATION ERROR: {str(e)}]")
                run.font.size = NORMAL_SIZE
                run.font.name = FONT_NAME
        else:
            run = paragraph.add_run(segment)
            run.font.superscript = is_superscript
            run.font.subscript = is_subscript
            run.font.size = NORMAL_SIZE
            run.font.name = FONT_NAME
